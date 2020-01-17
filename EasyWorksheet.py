import nltk
from tkinter import *
import tkinter as tk
from tkinter import filedialog
from docx import *
from collections import OrderedDict
from firebase import firebase
from PIL import Image
from google.cloud import storage
from urllib import request
import docx2txt
import pyrebase

ABERRANT_PLURAL_MAP = {
    'appendix': 'appendices',
    'barracks': 'barracks',
    'cactus': 'cacti',
    'child': 'children',
    'criterion': 'criteria',
    'deer': 'deer',
    'echo': 'echoes',
    'elf': 'elves',
    'embargo': 'embargoes',
    'focus': 'foci',
    'fungus': 'fungi',
    'goose': 'geese',
    'hero': 'heroes',
    'hoof': 'hooves',
    'index': 'indices',
    'knife': 'knives',
    'leaf': 'leaves',
    'life': 'lives',
    'man': 'men',
    'mouse': 'mice',
    'nucleus': 'nuclei',
    'person': 'people',
    'phenomenon': 'phenomena',
    'potato': 'potatoes',
    'self': 'selves',
    'syllabus': 'syllabi',
    'tomato': 'tomatoes',
    'torpedo': 'torpedoes',
    'veto': 'vetoes',
    'woman': 'women',
    }
VOWELS = set('aeiou')
nouns = []
convert_= []
lines = 0
top = 0
top2 = 0
top3 = 0 
top4 = 0
top5 = 0
top6 = 0
top7 = 0
top9 = 0
lines = ""
Selected_ = 0 
WordToReplace = ""
nouns_ = ""
FinalText = ""
splitText = []
len_split = 0
anti_nouns = []
login_entry = ""
pass_entry = ""
Error_label = ""
Success_Error = ""
Found = False
Real_Username = ""
input_ = ""
Name_for_Image = ""
Entry1 = ""
Entry2 = ""
Name_for_docx = ""
name_ = ""
Label_worksheet = ""
Button_worksheet = ""
Entry_worksheet = ""
Name_for_docx1 = ""
Label_notify = ""
v = ""

def pluralize(singular):
    
    if not singular:
        return ''
    plural = ABERRANT_PLURAL_MAP.get(singular)
    if plural:
        return plural
    root = singular
    try:
        if singular[-1] == 'y' and singular[-2] not in VOWELS:
            root = singular[:-1]
            suffix = 'ies'
        elif singular[-1] == 's':
            if singular[-2] in VOWELS:
                if singular[-3:] == 'ius':
                    root = singular[:-2]
                    suffix = 'i'
                else:
                    root = singular[:-1]
                    suffix = 'ses'
            else:
                suffix = 'es'
        elif singular[-2:] in ('ch', 'sh'):
            suffix = 'es'
        else:
            suffix = 's'
    except IndexError:
        suffix = 's'
    plural = root + suffix
    return plural

if __name__ == '__main__':
    import doctest
    doctest.testmod()

def noun_uni(l):
	global anti_nouns
	[anti_nouns.append(x) for x in l if x not in anti_nouns]
	return anti_nouns
    


def splitInput(lines):
	counter = 0
	global Selected_
	global lines_
	global FinalText
	global splitText
	global WordToReplace
	global firebase

	WordToReplace = Selected_ + ".png"

	config = {

  	"apiKey": "apiKey",
  	"authDomain": "AIzaSyCIbD3AEbT41u4UgudSa12gT7c4jADWZ6Y",
  	"databaseURL": "https://vaibhav-db.firebaseio.com/m",
  	"storageBucket": "vaibhav-db.appspot.com",
  	"serviceAccount": "mydb-1baf6a4b60a7.json"
	}

	
	abc = ["1"]

	for xsyaud in abc:
		firebase_img = pyrebase.initialize_app(config)
		storage_img = firebase_img.storage()
		try:
			storage_img.child(WordToReplace).download(WordToReplace)
		except:
			messagebox.showinfo("Error", "The noun identified is not in the database, please restart the program.")
			break
	


	lines_ = lines.split()
	full_stop = Selected_+"."

	for words in lines_ :
		

		if words.upper() == Selected_.upper() or words.upper() == full_stop.upper() :
			lines_[counter] = "______" #6 dashes
		counter = counter + 1

	FinalText = " ".join(lines_)
	splitText = FinalText.split("______")


	print (splitText)

	ToWord()

def ToWord():
	global len_split
	global FinalText
	global splitText
	global Name_for_docx
	global Real_Username
	global name_

	counter_split = 0
	name_ = Name_for_docx.get()
	name_final = name_ + ".docx"

	x = "user_files/" + Real_Username + "/" + name_final
	y = "/Files/" + Real_Username


	
	len_split = len(splitText) - 1
	document = Document()
	Paragraph1 = document.add_paragraph()
	v = Paragraph1.add_run()
	while counter_split < len_split:


		v.add_text (splitText[counter_split])
		v.add_text ('______ (')
		v.add_picture (WordToReplace)
		v.add_text (')')
		counter_split = counter_split + 1

	v.add_text(splitText[counter_split])



	
	document.save(name_final)

	config_ = {

  	"apiKey": "apiKey",
  	"authDomain": "AIzaSyCIbD3AEbT41u4UgudSa12gT7c4jADWZ6Y",
  	"databaseURL": "https://vaibhav-db.firebaseio.com/m",
  	"storageBucket": "vaibhav-db.appspot.com",
  	"serviceAccount": "mydb-1baf6a4b60a7.json"
	}

	firebase_word = pyrebase.initialize_app(config_)
	storage_img = firebase_word.storage()
	storage_img.child(x).put(name_final)

	#link_data = storage_img.child(x).get_url()
	Make_list()

def ToWord_GUEST():
	global len_split
	global FinalText
	global splitText
	

	counter_split = 0
	

	


	
	len_split = len(splitText) - 1
	document = Document()
	Paragraph1 = document.add_paragraph()
	v = Paragraph1.add_run()
	while counter_split < (len_split):


		v.add_text (splitText[counter_split])
		v.add_text ('______ (')
		v.add_picture (WordToReplace)
		v.add_text (')')
		counter_split = counter_split + 1

	v.add_text(splitText[counter_split])



	
	document.save("Converted.docx")
	
def Make_list():
	global firebase
	global name_
	global Real_Username
	x = Real_Username
	print (x)
	
	firebase_to_list = firebase.FirebaseApplication('https://vaibhav-db.firebaseio.com/', None)
	firebase_to_list.put( x , name_, "Image")

def getSelect(element_):
	global Selected_
	Selected_ = element_

	print (element_)
	splitInput(lines)
	#f = open('demo1.docx', 'rb')
	Label_Top3 = tk.Label (top2, text = 'The worksheet has been generated.'  ).place(x = 0 , y = 460)	
	exit_W2 = tk.Button (top2,text = "Exit", command = exitW2 ).place(x =0 , y = 490)

def getSelect1(element_):
	global Selected_
	Selected_ = element_

	print (element_)
	splitInput_GUEST(lines)
	#f = open('demo1.docx', 'rb')
	Label_Top3 = tk.Label (top7, text = 'The worksheet has been generated and saved as converted.docx!'  ).place(x = 0 , y = 460)	
	exit_W2 = tk.Button (top7,text = "Exit", command = exitW7 ).place(x =0 , y = 490)

def exitW1():
	top.destroy()

def exitW2():
	top2.destroy()

def exitW3():
	top3.destroy()

def exitW4():
	top4.destroy()

def exitW5():
	top5.destroy()

def exitW6():
	top6.destroy()

def exitW7():
	top7.destroy()


def exitW8():
	top8.destroy()


def exitW9():
	top9.destroy()




def getVal() :
	global lines
	global top
	global input_
	counter = 0
	lines = input_.get()
	
	is_noun = lambda pos: pos[:2] == 'NN'
	tokenized = nltk.word_tokenize(lines)
	global nouns
	nouns = [word for (word, pos) in nltk.pos_tag(tokenized) if is_noun(pos)] 
	
	noun_uni(nouns)
	

	for x in anti_nouns:
			convert_.append(x)
				
		

	print (convert_)
	exitW1()
	W7()

def getVal1() :
	global top
	global input_
	global lines
	counter = 0
	lines = input_.get()
	
	is_noun = lambda pos: pos[:2] == 'NN'
	tokenized = nltk.word_tokenize(lines)
	global nouns
	nouns = [word for (word, pos) in nltk.pos_tag(tokenized) if is_noun(pos)] 
	
	noun_uni(nouns)
	

	for x in anti_nouns:
			convert_.append(x)
				
		

	print (convert_)
	exitW6()
	W2()

def Check_Id():
	global Found
	global login_entry
	global pass_entry
	global firebase
	global Real_Username
	
	firebase1 = firebase.FirebaseApplication('https://vaibhav-db.firebaseio.com/', None)
	id_= firebase1.get('/Users' ,None)
	print (id_)
	a = login_entry.get()
	b = pass_entry.get()
	print (a)
	print (b)

	for key in id_:
		if key == a and id_[key] == b :
			Real_Username = a
			Found  = True
			exitW4()
			W6()

	if Found == False :
		Error_label.place(x = 0 , y = 470)

def Upload_database():
	global Val_img
	global Name_for_Image
	Val_img = Name_for_Image.get()
	c = Val_img + ".png"
	plural_ = pluralize(Val_img)
	cx = plural_ + ".png"
	storage_client = storage.Client.from_service_account_json('mydb-1baf6a4b60a7.json')
	buckets = storage_client.get_bucket('vaibhav-db.appspot.com')
	print(buckets)
	Selected_File = filedialog.askopenfilename()
	blob = buckets.blob(c)
	blob.upload_from_filename(filename= Selected_File)

	blob1 = buckets.blob(cx)
	blob1.upload_from_filename(filename= Selected_File)

	firebase4 = firebase.FirebaseApplication('https://vaibhav-db.firebaseio.com/', None)
	abcd = "Image"
	abcd12 = Val_img

	firebase4.put("/Images", abcd12, abcd12)

def func_upload() :
	global top6
	global Name_for_Image
	global Val_img
	Label_new = Label (top6, text = "Please type the name of the noun: ")
	Label_new.place (x = 0, y = 430)
	Name_for_Image.place(x = 235 , y = 430)
	Buttonxy = tk.Button (top6, text = "Select", command = Upload_database)
	Buttonxy.place(x = 430 , y = 435)
	
def Check_Done():
	global Entry1
	global Entry2
	global firebase
	global top5 
	global Label_Error
	Found = False

	c1 = Entry1.get()
	c2 = Entry2.get()
	firebase3 = firebase.FirebaseApplication('https://vaibhav-db.firebaseio.com/', None)
	data = firebase3.get('/Users' ,None)


	for z in data:
		if z == c1 :
			Label_Error.place (x = 0 , y = 430)
			Found = True

	if Found == False :
		firebase3.put("/Users",c1,c2)
		x = "/Files/" + c1
		firebase3.put(c1,"Demo","image")
		config_ = {

  		"apiKey": "apiKey",
  		"authDomain": "AIzaSyCIbD3AEbT41u4UgudSa12gT7c4jADWZ6Y",
  		"databaseURL": "https://vaibhav-db.firebaseio.com/m",
  		"storageBucket": "vaibhav-db.appspot.com",
  		"serviceAccount": "mydb-1baf6a4b60a7.json"
		}
		document = Document()
		Paragraph1 = document.add_paragraph()
		document.save("Converted.docx")
		firebase_word = pyrebase.initialize_app(config_)
		storage_img = firebase_word.storage()
		www = "user_files/" + c1 + "/demo.docx"
		storage_img.child(www).put("Converted.docx")
		x = c1
		print (x)
		Success_Error.place (x = 0 , y = 430)


def from_word() :
	global top2
	global top6
	global top8
	global convert_ 
	global lines
	global Selected_
	global lines_
	global FinalText
	global splitText
	global WordToReplace
	global firebase
	global len_split
	global Name_for_docx1
	global Real_Username
	global name_
	global v
	
	
	
	config = {

  	"apiKey": "apiKey",
  	"authDomain": "AIzaSyCIbD3AEbT41u4UgudSa12gT7c4jADWZ6Y",
  	"databaseURL": "https://vaibhav-db.firebaseio.com/m",
  	"storageBucket": "vaibhav-db.appspot.com",
  	"serviceAccount": "mydb-1baf6a4b60a7.json"
	}
	firebase_img = pyrebase.initialize_app(config)
	storage_img = firebase_img.storage()

	

	Selected_File1 = filedialog.askopenfilename()
	lines123= docx2txt.process(Selected_File1)
	lines1 = lines123.replace('\n', ' ').replace('\r', '')
	print (lines1)
	convert_all = lines1.split(".")

	print (convert_all)
	
	convert_all.pop()

  
	
	
	counter_final = 0
	
	for  lines in convert_all:
		print (lines)
		counter_split = 0
		counter = 0


		is_noun = lambda pos: pos[:2] == 'NN'
		tokenized = nltk.word_tokenize(lines)
		nouns = [word for (word, pos) in nltk.pos_tag(tokenized) if is_noun(pos)] 
	
		print (nouns)
		

		

		try:
			for nouns_ in nouns :
				print(nouns_)
				counter = 0
				img_ = nouns_ + ".png"
				storage_img.child(img_).download(img_)
		except:
			messagebox.showinfo("Error", "The noun identifies is not in the database, please restart the program.")
			break



		for bcde in nouns:
	 		if bcde == "i" or bcde == "I" :
	 			convert_.remove(bcde)




		len_nouns = len(nouns)
		print(len_nouns)

		if len_nouns == 1:
			nouns_ = nouns[0]

		
			lines_ = lines.split()
			print (lines_)
	
			full_stop = nouns_+"."


			for words in lines_ :

				print (words)
				
				if words.upper() == nouns_.upper() or words.upper() == full_stop.upper() :
					lines_[counter] = "______" #6 dashes
					
				counter = counter + 1

			FinalText = " ".join(lines_)
			splitText = FinalText.split("______")

			
			len_split = len(splitText) - 1
			
			if counter_final == 0 :

			
				name_ = Name_for_docx1.get()
				name_final = name_ + ".docx"

				x = "user_files/" + Real_Username + "/" + name_final
				y = "/Files/" + Real_Username
				document = Document()
				Paragraph1 = document.add_paragraph()
			
				v = Paragraph1.add_run()
			
				while counter_split < len_split:


					v.add_text (splitText[counter_split])
					v.add_text ('______ (')
					v.add_picture (img_)
					v.add_text ('). ')
					counter_split = counter_split + 1 
			else:
				while counter_split < len_split:


					v.add_text (splitText[counter_split])
					v.add_text ('______ (')
					v.add_picture (img_)
					v.add_text ('). ')
					counter_split = counter_split + 1 

		counter_final = counter_final + 1
	




	document.save(name_final)

	config_ = {

  	"apiKey": "apiKey",
  	"authDomain": "AIzaSyCIbD3AEbT41u4UgudSa12gT7c4jADWZ6Y",
  	"databaseURL": "https://vaibhav-db.firebaseio.com/m",
  	"storageBucket": "vaibhav-db.appspot.com",
  	"serviceAccount": "mydb-1baf6a4b60a7.json"
	}

	firebase_word = pyrebase.initialize_app(config_)
	storage_img = firebase_word.storage()
	storage_img.child(x).put(name_final)
	aasad = "Image"
	zx = Real_Username
	firebase5 = firebase.FirebaseApplication('https://vaibhav-db.firebaseio.com/', None)
	firebase5.put(zx, name_, aasad)

	Name_for_docx_Label1 = Label (top8, text= "The worksheet has been generated." ).place(x = 0, y = 360)
	Button_exit = tk.Button (top8, text = "Exit" , command = exitW8).place(x =0 , y = 400)


def logout() :
	exitW6()
	
	W3()
	
def back_():
	exitW5()
	W3()

def back1_():
	exitW4()
	W3()

def back_home():
	exitW1()
	W3()

def Download_file():
	global Entry_worksheet
	global Real_Username
	global Label_notify

	xyzz = Entry_worksheet.get()

	config = {

  	"apiKey": "apiKey",
  	"authDomain": "AIzaSyCIbD3AEbT41u4UgudSa12gT7c4jADWZ6Y",
  	"databaseURL": "https://vaibhav-db.firebaseio.com/m",
  	"storageBucket": "vaibhav-db.appspot.com",
  	"serviceAccount": "mydb-1baf6a4b60a7.json"
	}




	firebase_img = pyrebase.initialize_app(config)
	storage_img = firebase_img.storage()
	storage_img.child("user_files").child(Real_Username).child(xyzz).download(xyzz)

	Label_notify.place(x= 0 , y = 390)
	


def word123() :
	exitW6()
	W8()

def back_W6():
	exitW9()
	W6()


def place_stuff():
	global Label_worksheet 
	global Button_worksheet 
	global Entry_worksheet

	Label_worksheet.place(x = 0 , y = 550)
	Entry_worksheet.place(x = 150 , y = 550 )
	Button_worksheet.place(x = 0, y = 590)

def W1():
	exitW3()
	global top
	global input_
	top = tk.Tk()
	logo = tk.PhotoImage(file="JPIS.gif")
	SHOW1 = tk.Label(top, image=logo).pack(side="top")
	Label1 = Label (top, text="sentence you want to convert: ")
	Label1.place(x = 0 , y = 310)
	input_ = Entry(top)
	input_.place(x = 200 , y = 310)
	convertB = tk.Button (top,text = "convert", command = getVal )
	convertB.place(x = 390 , y = 310)
	
	xy = tk.Button(top, text = "Back" , command = back_home ).place(x = 0 , y = 400)

	top.geometry ( "500x500+350+150" )
	top.title ("Convert Vocabulary 101")

	top.mainloop()

def W2(): 
	global top2
	global Name_for_docx


	
	top2 = tk.Tk()
	logo2 = tk.PhotoImage(file="JPIS.gif")
	SHOW2= tk.Label(top2, image=logo2).pack(side="top")

	Name_for_docx_Label = Label (top2, text= "Please enter the name for the worksheet: " )
	Name_for_docx = Entry(top2)
	Name_for_docx_Label.place ( x= 0 , y = 310)
	Name_for_docx.place( x = 300 , y = 310 )
	




	Label_Top2 = tk.Label (top2, text = "These are the identified nouns, please select one of the following to convert:" ).place(x = 0, y = 350)

	y = 0
	place = 370
	for element_ in convert_ : 
		
		if y > 600 :
			place = place + 40
			y = 0

		tk.Button (top2, text = element_, command = lambda s=element_: getSelect(s)).place( x = y , y = place)
		y = y + 80

	

	
	top2.geometry ( "680x600+350+150" )
	top2.title ("Convert Vocabulary 101")
	top2.mainloop()

def W3():
	global top3
	top3 = tk.Tk()
	logo3 = tk.PhotoImage(file="JPIS.gif")
	SHOW3 = tk.Label(top3, image=logo3).pack(side="top")
	Teacher_label = Label (top3, text="Are you a existing user: ")
	Teacher_Button = tk.Button (top3, text = "Log in!" , command = W4)
	Guest_label = Label (top3, text="Sign up: ")
	Guest_Button = tk.Button (top3, text = "Create" , command = W5)
	Teacher_label.place(x = 0 , y = 310)
	Teacher_Button.place(x = 150 , y = 310)
	Guest_label.place(x = 0 , y = 350)
	Guest_Button.place(x = 60 , y = 350)
	New_Button = tk.Button (top3, text = "Exit" , command = exitW3 )
	New_Button.place(x = 0 , y = 430)
	top3.geometry ( "500x500+350+150" )
	top3.title ("WORKSHEET GENERATOR")
	top3.mainloop()

def W4():
	global top4
	exitW3()
	global login_entry
	global pass_entry
	global Error_label
	top4 = tk.Tk()
	logo4 = tk.PhotoImage(file="JPIS.gif")
	SHOW3 = tk.Label(top4, image=logo4).pack(side="top")
	login_label = Label (top4, text="Please enter your Username: ")
	login_entry= Entry(top4)
	pass_label = Label (top4, text="Please enter your Password: ")
	pass_entry= Entry(top4, show = "*") 
	main_Button = tk.Button (top4, text = "Sign In!" , command = Check_Id )
	Error_label = Label (top4, text = "The given credentials are incorrect, please try again.")
	login_label.place(x = 0 , y = 310)
	login_entry.place(x = 200 , y = 310)
	pass_label.place(x = 0 , y = 350)
	pass_entry.place(x = 200 , y = 350)
	main_Button.place(x = 0 , y = 390)
	New_Button2 = tk.Button (top4, text = "Back" , command = back1_)
	New_Button2.place(x = 0 , y = 430)

	top4.geometry ( "500x500+350+150" )
	top4.title ("Convert Vocabulary 101")
	top4.mainloop()

def W5():
	exitW3()
	global top5
	global Entry1
	global Entry2
	global Label_Error
	global Success_Error
	top5 = tk.Tk()
	logo5 = tk.PhotoImage(file="JPIS.gif")
	SHOW5 = tk.Label(top5, image=logo5).pack(side="top")
	Label_add = Label (top5, text = "Please choose a username:")
	Label_add.place(x = 0 , y = 310)
	Entry1 = Entry (top5)
	Entry1.place (x = 200 , y = 310)
	Label2_add = Label (top5, text = "Please choose a Password:")
	Label2_add.place(x = 0 , y = 350)
	Entry2 = Entry (top5, show = "*")
	Entry2.place (x = 200 , y = 350)
	upload_ = tk.Button (top5 ,text = "Sign Up", command = Check_Done)
	upload_.place(x=0 , y= 390)
	Label_Error = Label (top5, text = " Username already taken, please choose another one. ")
	Success_Error = Label (top5, text = " Your ID has been created. ")
	New_Button2 = tk.Button (top5, text = "Back" , command = back_)
	New_Button2.place(x = 0 , y = 470)
	top5.geometry ( "500x500+350+150" )
	top5.mainloop()

def W6():
	global top9
	global Real_Username
	global input_
	global Name_for_Image
	global top6
	global firebase
	global Label_worksheet 
	global Button_worksheet 
	global Entry_worksheet

	
	
	Real_Username1 = "Welcome " + Real_Username + "!"
	top6 = tk.Tk()
	logo6 = tk.PhotoImage(file="JPIS.gif")
	SHOW6 = tk.Label(top6, image=logo6).pack(side="top")
	Label_Username = Label (top6, text = Real_Username1)
	Label_Username.place(x = 200 , y = 310)
	Label1 = Label (top6, text="Input sentence(multiple nouns): ")
	Label1.place(x = 0 , y = 350)
	input_ = Entry(top6)
	input_.place(x = 220 , y = 350)
	convertB = tk.Button (top6,text = "convert", command = getVal1 )
	convertB.place(x = 420 , y = 355)
	Labelxx = Label (top6, text= "Insert an Image into the database:" )
	Add_img = tk.Button (top6,text = "Add" , command = func_upload)
	Name_for_Image = Entry(top6)
	Labelxx.place(x = 0 , y = 390)
	Add_img.place(x = 230 , y = 390)
	word_Label = Label (top6, text = "Upload an existing word document: ")
	word_button = tk.Button (top6, text = "open" , command = word123 )
	word_Label.place(x=0 , y = 470)
	word_button.place( x = 240 , y = 470)
	New_Button1 = tk.Button (top6, text = "Log 0ut" , command = logout )
	New_Button1.place(x = 0 , y = 550 )
	
	x = Real_Username
	firebase_x = firebase.FirebaseApplication('https://vaibhav-db.firebaseio.com/', None)
	
	
	xyz = tk.Button(top6, text = "Download a previous project?" , command = W9 ).place(x = 0, y = 510)
	
	
	data2 = firebase_x.get("Images",None)
	List_box2 = Listbox(top6)


	numbers_ = 1

	for f in data2:
		List_box2.insert(numbers_,f)
		numbers_ = numbers_ + 1

	List_box2.place(x = 550, y = 320)
	
	Scrollbar(List_box2, orient = "vertical")

	dsad = Label(top6, text = "Nouns in the Database: ").place(x = 550, y = 290)


	print (data2)






	top6.geometry ( "740x600+350+10" )
	top6.title ("Convert Vocabulary 101")

	top6.mainloop()

def W7():
	global top7
	global Name_for_docx


	
	top7 = tk.Tk()
	logo7 = tk.PhotoImage(file="JPIS.gif")
	SHOW7= tk.Label(top7, image=logo7).pack(side="top")

	




	Label_Top2 = tk.Label (top7, text = "These are the identified nouns, please select one of the following to convert:" ).place(x = 0, y = 310)

	y = 0
	xy = 330


	for element_ in convert_ : 
		

		tk.Button (top7, text = element_, command = lambda s=element_: getSelect1(s)).place( x = y , y = xy)
		y = y + 80

		if y > 590 :

			xy = xy + 40

	

	

	
	top7.geometry ( "680x600+350+150" )
	top7.title ("Convert Vocabulary 101")
	top7.mainloop()

def W8(): 
	global top8
	global Name_for_docx1

	


	
	top8 = tk.Tk()
	logo8 = tk.PhotoImage(file="JPIS.gif")
	SHOW8= tk.Label(top8, image=logo8).pack(side="top")

	Name_for_docx_Label1 = Label (top8, text= "Please enter the name for the worksheet: " )
	Name_for_docx1 = Entry(top8)
	Name_for_docx_Label1.place ( x= 0 , y = 310)
	Name_for_docx1.place( x = 300 , y = 310 )
	New_Button1 = tk.Button (top8, text = "Convert" , command = from_word )

	New_Button1.place(x = 500, y = 310 )

	top8.geometry ( "680x600+350+150" )
	top8.title ("Worksheet Generator")
	top8.mainloop()

def W9():
	exitW6()
	global top9
	global Real_Username
	global Entry_worksheet
	global Label_notify

	top9 = tk.Tk()
	logo9 = tk.PhotoImage(file= "JPIS.gif")
	SHOW9 = tk.Label(top9, image=logo9).pack(side="top")

	x = Real_Username
	firebase_y = firebase.FirebaseApplication('https://vaibhav-db.firebaseio.com/', None)
	data = firebase_y.get(x ,None)
	List_box1 = Listbox(top9)

	y = 1
	
	for z in data:
		z =  z + ".docx"
		List_box1.insert(y,z)
		y = y + 1

	Label_worksheet1  = Label(top9, text = " Name the worksheet: " )
	Entry_worksheet = Entry(top9)
	Button_worksheet = tk.Button(top9, text = "Download" , command = Download_file )

	Label_notify  = Label(top9, text = "The worksheet has been downloaded" )


	xx = Label(top9, text = " Your previous worksheets :").place(x = 350 , y = 290)

	Label_worksheet1.place(x= 0 , y = 310)
	Entry_worksheet.place(x= 150 , y = 310)
	Button_worksheet.place(x= 0 , y = 350)
	List_box1.place(x = 350, y = 320)

	Button_back = tk.Button(top9, text = "Back" , command = back_W6 ).place(x=0, y = 500)

	top9.geometry ( "580x540+350+150" )
	top9.title ("Worksheet Generator")
	top9.mainloop()












W3()












 

 
